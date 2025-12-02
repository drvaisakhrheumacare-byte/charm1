import streamlit as st
import pandas as pd
import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import secrets
import string
import io
import re
import datetime

# --- CONFIGURATION ---
DEFAULT_SHEET_URL = "https://docs.google.com/spreadsheets/d/1mYapaNzFhSdWTLWaK1cedvQVcjCgvr17EQ-SkEpwV24/edit?gid=0#gid=0"
COUPON_VALUE = "₹10"

# --- HELPER FUNCTIONS ---

def get_sheet_csv_url(original_url):
    """Extracts ID and GID (Tab ID) to create a CSV export link."""
    if not original_url:
        return None, "Empty URL"
    
    try:
        # Extract Sheet ID
        sheet_id_match = re.search(r'/d/([a-zA-Z0-9-_]+)', original_url)
        if not sheet_id_match:
            return None, "Could not find Sheet ID."
        sheet_id = sheet_id_match.group(1)

        # Extract GID (The specific tab ID)
        gid = "0"
        gid_match = re.search(r'[#&?]gid=([0-9]+)', original_url)
        if gid_match:
            gid = gid_match.group(1)
            
        return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}", None
    except Exception:
        return None, "Invalid URL format."

def generate_secure_code(prefix):
    """Generates unique code: PREFIX-XXX-XXXX"""
    alphabet = string.ascii_uppercase + string.digits
    part1 = ''.join(secrets.choice(alphabet) for _ in range(3))
    part2 = ''.join(secrets.choice(alphabet) for _ in range(4))
    
    clean_prefix = str(prefix).upper().strip()
    
    # Check for invalid prefixes
    invalid_prefixes = ['NAN', 'NONE', '', 'nan']
    if clean_prefix in invalid_prefixes:
        return f"{part1}-{part2}"
        
    return f"{clean_prefix}-{part1}-{part2}"

def create_coupon_content(cell, name, emp_id, code, date_label):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # 1. Value
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(f"{COUPON_VALUE} Coupon")
    run1.bold = True
    run1.font.size = Pt(9)
    run1.font.color.rgb = RGBColor(0, 100, 0)
    
    # 2. Unique Code
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(code)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(10)
    run2.bold = True
    
    # 3. Employee Info
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"{name}\n({emp_id})")
    run3.font.size = Pt(7)
    
    # 4. Month/Year
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(date_label)
    run4.font.size = Pt(6)
    run4.font.italic = True

def generate_docx(df, selected_date, default_prefix_input):
    doc = docx.Document()
    
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # --- INTELLIGENT COLUMN DETECTION ---
    col_name = df.columns[0]
    col_id = df.columns[1]
    col_prefix = None

    for col in df.columns:
        if "name" in col.lower():
            col_name = col
        if "code" in col.lower() or "id" in col.lower():
            col_id = col
        if "prefix" in col.lower():
            col_prefix = col
    
    st.info(f"Using columns: Name='{col_name}', ID='{col_id}', Prefix='{col_prefix if col_prefix else 'Using Default'}'")

    progress_bar = st.progress(0)
    total_rows = len(df)

    for index, row in df.iterrows():
        emp_name = str(row[col_name])
        emp_id = str(row[col_id])
        
        # --- DETERMINE PREFIX ---
        current_prefix = default_prefix_input
        if col_prefix:
            row_prefix = str(row[col_prefix])
            if row_prefix.lower() != 'nan' and row_prefix.strip():
                current_prefix = row_prefix

        # --- HEADER (Top of Page) ---
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Removed Prefix from here as requested
        run_h = header_p.add_run(f"Name: {emp_name} ({emp_id})  |  Month: {selected_date}")
        run_h.bold = True
        run_h.font.size = Pt(12)
        run_h.font.name = 'Arial'
        header_p.paragraph_format.space_after = Pt(6)

        # --- TABLE (Coupons) ---
        table = doc.add_table(rows=13, cols=5)
        table.style = 'Table Grid'
        
        for row_obj in table.rows:
            row_obj.height = Cm(1.9) 
            
        for r in range(13):
            for c in range(5):
                cell = table.cell(r, c)
                unique_code = generate_secure_code(current_prefix)
                create_coupon_content(cell, emp_name, emp_id, unique_code, selected_date)
        
        # New Page for next employee
        if index < len(df) - 1:
            doc.add_page_break()
        
        progress_bar.progress((index + 1) / total_rows)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- MAIN APP UI ---
st.set_page_config(page_title="Coupon Generator", page_icon="🎫")

# Initialize Session State
if "password_correct" not in st.session_state:
    st.session_state["password_correct"] = False

def check_password():
    if st.session_state["password_correct"]:
        return True
    
    if "password" not in st.secrets:
        return True

    pwd = st.text_input("Enter App Password", type="password")
    if pwd == st.secrets["password"]:
        st.session_state["password_correct"] = True
        return True
    elif pwd:
        st.error("Incorrect Password")
    return False

if check_password():
    st.title("🎫 Monthly Coupon Generator")

    with st.container():
        st.subheader("1. Select Month")
        now = datetime.datetime.now()
        current_year = now.year
        months = ["January", "February", "March", "April", "May", "June", 
                  "July", "August", "September", "October", "November", "December"]
        years = [current_year + i for i in range(11)] 
        
        c1, c2 = st.columns(2)
        sel_month = c1.selectbox("Month", options=months, index=now.month - 1)
        sel_year = c2.selectbox("Year", options=years, index=0)
        selected_date = f"{sel_month} {sel_year}"

        st.subheader("2. Settings")
        sheet_url = st.text_input("Google Sheet URL", value=DEFAULT_SHEET_URL)
        default_prefix = st.text_input("Fallback Prefix (Used if 'Prefix' column is empty)", value="EMP")

    if st.button("Generate Coupons", type="primary"):
        st.write("---")
        with st.spinner('Fetching Data...'):
            try:
                # 1. Fetch Sheet
                link, err = get_sheet_csv_url(sheet_url)
                if err:
                    st.error(f"Error: {err}")
                    st.stop()
                
                df = pd.read_csv(link)
                df.columns = [c.strip() for c in df.columns]
                st.write(f"✅ Loaded {len(df)} Employees.")

                # 2. Generate
                docx_file = generate_docx(df, selected_date, default_prefix)
                
                safe_date = re.sub(r'[^\w\s-]', '', selected_date).strip().replace(' ', '_')
                file_name = f"Coupons_{safe_date}.docx"

                st.success("🎉 Generation Complete!")
                st.download_button(
                    label="📥 Download File",
                    data=docx_file,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                    
            except Exception as e:
                st.error(f"An error occurred: {e}")
